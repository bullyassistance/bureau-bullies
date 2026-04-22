"""
Bully AI — Report + Attack Plan .docx generator
------------------------------------------------
After a scan, we generate a Word doc the consumer can download:
  - Their scan results (leverage, violations, case law)
  - A custom attack plan (week-by-week)
  - Three pre-filled dispute letters (one per bureau)
  - One validation letter + one demand letter (ready to customize)

Uses python-docx — no external dependencies.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

from analyzer import ReportSummary

RED = RGBColor(0xE1, 0x1D, 0x2E)
DARK = RGBColor(0x10, 0x11, 0x18)
MUTE = RGBColor(0x6B, 0x72, 0x80)


def _heading(doc, text, size=20, color=RED, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return p


def _sub(doc, text, size=11, color=MUTE, italic=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.font.color.rgb = color
    return p


def _body(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def _rule(doc):
    doc.add_paragraph("_" * 70).runs[0].font.color.rgb = MUTE


def generate_report_doc(
    summary: ReportSummary,
    consumer_first: str,
    consumer_last: str,
    consumer_email: str,
    consumer_phone: str,
    out_path: Path,
) -> Path:
    """Generate the full client deliverable — results + attack plan + letters."""
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Cover ---------------------------------------------------------
    _heading(doc, "THE BUREAU BULLIES", size=28, color=RED, bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sub(doc, "Bully AI — Deep-Dive Credit Intelligence Report", size=13, color=DARK, italic=False).alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sub(doc, f"Prepared for: {consumer_first} {consumer_last}  ·  {date.today().isoformat()}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    _rule(doc)

    # ---- Executive summary --------------------------------------------
    _heading(doc, "EXECUTIVE SUMMARY", size=16)
    _body(doc, summary.executive_summary or "Report analysis complete.")

    doc.add_paragraph()
    # Quick stats table
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Light Grid Accent 1"
    rows = [
        ("Total Estimated Leverage", f"${summary.total_estimated_leverage:,.2f}"),
        ("Violations Identified",    str(len(summary.violations))),
        ("Negative Items",           str(summary.total_negative_items)),
    ]
    for i, (k, v) in enumerate(rows):
        tbl.cell(i, 0).text = k
        tbl.cell(i, 1).text = v

    doc.add_paragraph()
    _body(doc, f"Top pain point: {summary.top_pain_point}")
    _body(doc, f"Recommended tier: {summary.recommended_tier.upper()}")
    _rule(doc)

    # ---- Violation list ------------------------------------------------
    _heading(doc, "VIOLATIONS IDENTIFIED", size=16)
    for i, v in enumerate(summary.violations, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. {v.creditor} (...{v.account_last4})  —  {v.bureau}")
        r.font.size = Pt(12); r.bold = True; r.font.color.rgb = RED
        _body(doc, f"Violation: {v.violation_type}")
        _body(doc, f"Details: {v.description}")
        if v.case_law_citation:
            _sub(doc, f"Backed by: {v.case_law_citation}", italic=True, color=MUTE)
        p = doc.add_paragraph()
        r = p.add_run(f"Estimated leverage: ${v.dollar_leverage:,.2f}  ·  Severity: {v.severity.upper()}")
        r.bold = True; r.font.color.rgb = DARK
        doc.add_paragraph()

    _rule(doc)

    # ---- The attack plan ----------------------------------------------
    _heading(doc, "YOUR ATTACK PLAN", size=16)
    plan = [
        ("WEEK 1 — Intelligence",
         "Pull fresh reports from all 3 bureaus (AnnualCreditReport.com). Screenshot every negative item (before-photos). Use the spreadsheet tracker to log each item."),
        ("WEEK 2 — Opening Volley",
         "Send the Debt Validation Letter (below) via CERTIFIED MAIL, return receipt, to the top collector. At the same time, file online disputes at all 3 bureaus using the bureau-specific letters below. Save every confirmation number."),
        ("WEEKS 3-4 — Wait & Watch",
         "Furnisher has 30 days (45 if you uploaded PDFs) to investigate. Pull your reports again at Day 31. Compare before/after."),
        ("WEEKS 5-6 — Escalate",
         "If items come back 'verified' with no real change, send the Demand Letter (below) via certified mail. Cite Hinkle v. Midland. 30-day clock starts on delivery."),
        ("WEEKS 7-10 — Enforce",
         "If they don't comply after the demand letter, you file a pro se federal complaint or engage The Bureau Bullies DFY team to file for you."),
    ]
    for title, body in plan:
        p = doc.add_paragraph()
        r = p.add_run(title); r.bold = True; r.font.color.rgb = RED; r.font.size = Pt(12)
        _body(doc, body)

    doc.add_page_break()

    # ---- LETTER 1: DEBT VALIDATION ------------------------------------
    _heading(doc, "LETTER 1 — DEBT VALIDATION (FDCPA § 1692g)", size=14)
    _sub(doc, "Send via certified mail, return receipt, within 30 days of first collector contact.")
    _body(doc,
        f"{consumer_first} {consumer_last}\n"
        f"[Your Mailing Address]\n"
        f"[City, State ZIP]\n"
        f"{date.today().isoformat()}\n\n"
        f"VIA CERTIFIED MAIL, RETURN RECEIPT REQUESTED\n\n"
        f"[COLLECTOR LEGAL NAME]\n"
        f"[Registered Agent / Legal Dept Address]\n"
        f"[City, State ZIP]\n\n"
        f"Re: Debt Validation Request — Account/Reference: {summary.top_collection_name or '[Account #]'}\n\n"
        f"To Whom It May Concern,\n\n"
        f"This letter is in response to your recent communication regarding the above-referenced "
        f"alleged debt. Pursuant to the Fair Debt Collection Practices Act, 15 U.S.C. § 1692g, "
        f"I am formally disputing this debt and requesting full validation.\n\n"
        f"I demand that you provide the following documentation within 30 days of receipt:\n"
        f"  1. Verification that the debt is owed and the amount is correct.\n"
        f"  2. A copy of the original signed contract or agreement.\n"
        f"  3. The name and address of the original creditor.\n"
        f"  4. Proof that your company is the current legal owner of this debt, including any "
        f"assignments or chains of title.\n"
        f"  5. Proof that your company is licensed to collect debts in my state.\n"
        f"  6. A complete account history including all payments, charges, and any interest/fees.\n\n"
        f"Until you provide this validation, you are required by law to cease collection activity. "
        f"Any further reporting to credit bureaus without proper validation may constitute a "
        f"violation of 15 U.S.C. § 1681s-2 and will be documented.\n\n"
        f"Respectfully,\n\n"
        f"{consumer_first} {consumer_last}\n"
        f"{consumer_phone}  ·  {consumer_email}\n"
    )
    doc.add_page_break()

    # ---- LETTER 2: CRA DISPUTE ----------------------------------------
    _heading(doc, "LETTER 2 — CREDIT BUREAU DISPUTE (15 U.S.C. § 1681s-2(b))", size=14)
    _sub(doc, "File this online at each bureau. Or send via certified mail to the addresses below.")
    _body(doc,
        "Bureau Addresses:\n"
        "  Experian:   P.O. Box 4500, Allen, TX 75013  ·  experian.com/disputes\n"
        "  Equifax:    Equifax Information Services LLC, P.O. Box 740256, Atlanta, GA 30374\n"
        "  TransUnion: TransUnion Consumer Solutions, P.O. Box 2000, Chester, PA 19016\n"
    )
    _body(doc,
        f"{consumer_first} {consumer_last}\n"
        f"[Your Mailing Address]\n"
        f"[City, State ZIP]\n"
        f"{date.today().isoformat()}\n\n"
        f"[BUREAU NAME AND ADDRESS]\n\n"
        f"Re: Formal Dispute of Inaccurate Information — 15 U.S.C. § 1681s-2(b)\n"
        f"Account/Reference Number: [ACCOUNT NUMBER]\n\n"
        f"To Whom It May Concern,\n\n"
        f"I am writing to dispute inaccurate information on my credit report maintained by your "
        f"agency. The following account contains information that is inaccurate and/or unverifiable:\n\n"
        f"Creditor/Furnisher: {summary.top_collection_name or '[CREDITOR]'}\n"
        f"Account Ending: [LAST 4]\n"
        f"Amount Reported: ${summary.top_collection_amount:,.2f}\n\n"
        f"Pursuant to 15 U.S.C. § 1681s-2(b)(1)(A)-(E), upon receipt of this dispute your agency "
        f"is required to notify the furnisher, who must then:\n"
        f"  (A) Conduct a reasonable investigation;\n"
        f"  (B) Review ALL relevant information I provide;\n"
        f"  (C) Report results back to your agency;\n"
        f"  (D) Report corrections to all nationwide CRAs if inaccurate;\n"
        f"  (E) Modify, delete, or permanently block information that CANNOT BE VERIFIED.\n\n"
        f"Per Hinkle v. Midland Credit Management, 800 F.3d 1295 (11th Cir. 2015), a furnisher's "
        f"internal database match does NOT constitute a reasonable investigation. The furnisher "
        f"must conduct a 'fairly searching inquiry' and produce original account-level documentation.\n\n"
        f"Grounds for Dispute:\n"
        f"  [Customize based on your situation — e.g., account is not mine, amount is incorrect, "
        f"payment history is wrong, account is beyond statute of limitations, etc.]\n\n"
        f"I request that you conduct your investigation, notify the furnisher, and report results "
        f"back to me within the 30-day period required by law.\n\n"
        f"Respectfully,\n\n"
        f"{consumer_first} {consumer_last}\n"
        f"{consumer_phone}  ·  {consumer_email}\n"
    )
    doc.add_page_break()

    # ---- LETTER 3: DEMAND LETTER --------------------------------------
    _heading(doc, "LETTER 3 — DEMAND LETTER (after 'verified' with no change)", size=14)
    _sub(doc, "Send AFTER your CRA dispute comes back 'verified' but nothing changed on your report. Certified mail, return receipt. 30-day clock.")
    _body(doc,
        f"{consumer_first} {consumer_last}\n"
        f"[Your Mailing Address]\n"
        f"[City, State ZIP]\n"
        f"{date.today().isoformat()}\n\n"
        f"VIA CERTIFIED MAIL, RETURN RECEIPT REQUESTED\n"
        f"Tracking #: _____________\n\n"
        f"[FURNISHER LEGAL NAME]\n"
        f"[Registered Agent / Legal Dept Address]\n\n"
        f"Re: Demand for Deletion — Violation of 15 U.S.C. § 1681s-2(b)\n"
        f"Account/Reference Number: [ACCOUNT #]\n\n"
        f"To Whom It May Concern,\n\n"
        f"This letter serves as formal legal demand regarding your violation of the Fair Credit "
        f"Reporting Act, specifically 15 U.S.C. § 1681s-2(b).\n\n"
        f"1. On [DATE], I filed a formal dispute with [BUREAU] regarding the above account. "
        f"Confirmation: [CRA CONFIRMATION #].\n"
        f"2. The credit bureau forwarded notice of my dispute to your company, triggering your "
        f"mandatory duties under § 1681s-2(b)(1)(A)-(E).\n"
        f"3. On [DATE], I received notice that your company 'verified' the account. My updated "
        f"credit report reflects NO meaningful change.\n\n"
        f"YOUR VIOLATIONS:\n"
        f"  [ ] Failed to conduct a reasonable investigation under § 1681s-2(b)(1)(A). Per Hinkle "
        f"v. Midland Credit Management, 800 F.3d 1295 (11th Cir. 2015), matching internal records "
        f"does not constitute a reasonable investigation.\n"
        f"  [ ] Failed to review documentation I submitted with my dispute under § 1681s-2(b)(1)(B).\n"
        f"  [ ] Failed to delete unverifiable information under § 1681s-2(b)(1)(E).\n"
        f"  [ ] Failed to report corrections to all nationwide CRAs under § 1681s-2(b)(1)(D).\n\n"
        f"DEMAND: Within thirty (30) days of your receipt of this letter, I demand that you:\n"
        f"  1. Permanently delete the above account from all consumer reporting agencies.\n"
        f"  2. Provide written confirmation of deletion.\n"
        f"  3. Immediately cease all collection activity.\n\n"
        f"If you fail to comply, I intend to file a civil action in United States District Court "
        f"seeking:\n"
        f"  - Statutory damages ($100–$1,000 per violation) under § 1681n.\n"
        f"  - Actual damages including emotional distress (Saunders v. BB&T, 526 F.3d 142, 4th Cir. 2008).\n"
        f"  - Punitive damages for willful noncompliance.\n"
        f"  - Attorney fees and costs.\n\n"
        f"I have preserved all documentation. This matter can be resolved without litigation if you "
        f"act within the timeframe stated above.\n\n"
        f"Respectfully,\n\n"
        f"{consumer_first} {consumer_last}\n"
        f"{consumer_phone}  ·  {consumer_email}\n"
    )

    # ---- Footer disclaimer --------------------------------------------
    doc.add_page_break()
    _heading(doc, "DISCLAIMER", size=12, color=MUTE)
    _body(doc,
        "The Bureau Bullies, LLC is a credit education and dispute services company located in "
        "Wilmington, Delaware. We are NOT a law firm and NOT attorneys. Nothing in this document "
        "constitutes legal advice, and no attorney-client relationship is created by receiving "
        "or using this report. This document is for educational and informational purposes only. "
        "Laws change, court interpretations evolve, and individual circumstances vary. Consult a "
        "licensed attorney in your jurisdiction before taking any legal action. Results are not "
        "guaranteed and individual outcomes vary."
    )

    doc.save(out_path)
    return out_path
